from __future__ import annotations

import csv
import json
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
DOWNLOADS = ROOT / "public" / "downloads" / "templates"
PREVIEW_OUTPUT = ROOT / "tools" / "template-artifacts" / "output"

BRAND = {
    "name": "Clarpoint",
    "site": "clarpoint.co",
    "tagline": "Clear plans. Stronger communication. Better execution.",
}

FORMAT_DIRS = {
    ".pdf": "pdf",
    ".docx": "word",
    ".xlsx": "excel",
    ".pptx": "powerpoint",
    ".csv": "csv",
}


PRODUCTS = {
    "executive-project-status-pack": {
        "title": "Executive Project Status Pack",
        "audience": "Project managers, delivery leads, consultants, PMO leads, and leadership-facing operators.",
        "summary": "A premium reporting and project visibility toolkit for teams that need cleaner executive updates, stronger ownership, and a sharper control layer.",
        "value_props": [
            "Executive-ready updates instead of scattered status notes",
            "Editable documents, dashboards, decks, and import-friendly trackers",
            "Sample content and blank working versions for same-day client or leadership use",
        ],
        "use_case": "Use this pack when active delivery work needs a clearer leadership narrative, visible risks, and stronger next-step accountability.",
        "preview_images": ["executive-project-control-workbook.png", "executive-deck.png"],
        "csv_defs": {
            "risk-and-escalation-register": [
                ["ID", "Risk / Escalation", "Impact", "Owner", "Mitigation", "Escalate By", "Status"],
                ["R-01", "Approval timing may affect milestone confidence", "Medium", "Program Manager", "Confirm decision path this week", "2026-06-02", "Open"],
                ["E-01", "Content dependency has not been resolved", "High", "Delivery Lead", "Raise sponsor intervention request", "2026-05-30", "In Progress"],
            ],
            "action-ownership-tracker": [
                ["Action ID", "Action", "Owner", "Due Date", "Priority", "Status", "Notes"],
                ["ACT-01", "Confirm milestone owner list", "Program Manager", "2026-05-29", "High", "In Progress", "Review in weekly readout"],
                ["ACT-02", "Draft leadership decision note", "Delivery Lead", "2026-05-31", "High", "Open", "Needs sponsor context"],
            ],
            "decision-log": [
                ["Decision ID", "Decision", "Decision Owner", "Date", "Impact", "Follow-Up"],
                ["DEC-01", "Keep phased launch approach", "Executive Sponsor", "2026-05-24", "Scope / timeline", "Refresh milestone plan"],
                ["DEC-02", "Escalate client approval dependency", "Program Manager", "2026-05-24", "Stakeholder communication", "Add to sponsor update"],
            ],
        },
    },
    "raid-log-action-tracker-bundle": {
        "title": "RAID Log + Action Tracker Bundle",
        "audience": "Project leads, operators, small delivery teams, account managers, and follow-up owners.",
        "summary": "A cleaner control bundle for teams that need one visible place for risks, issues, decisions, and next actions.",
        "value_props": [
            "Control files that reduce email-driven follow-up",
            "Blank and example trackers for faster adoption",
            "A practical review rhythm instead of a heavy PMO process",
        ],
        "use_case": "Use this bundle when work is moving but the team has no reliable source of truth for decisions, actions, and delivery pressure.",
        "preview_images": ["raid-action-control-workbook.png", "raid-deck.png"],
        "csv_defs": {
            "raid-log": [
                ["ID", "Category", "Item", "Impact", "Owner", "Mitigation", "Review Date", "Status"],
                ["R-01", "Risk", "Requirements are still shifting", "High", "Project Lead", "Freeze scope by next checkpoint", "2026-05-29", "Open"],
                ["I-01", "Issue", "Decision backlog is slowing approvals", "High", "Sponsor", "Prioritize open approvals in weekly review", "2026-05-30", "In Progress"],
            ],
            "action-tracker": [
                ["Action ID", "Action", "Owner", "Due Date", "Status", "Notes"],
                ["ACT-01", "Close outstanding scope questions", "Project Lead", "2026-05-28", "Open", "Needs client input"],
                ["ACT-02", "Update weekly risk summary", "PM", "2026-05-29", "In Progress", "Prepare for Friday review"],
            ],
            "decision-log": [
                ["Decision ID", "Decision", "Decision Owner", "Date", "Impact", "Next Step"],
                ["DEC-01", "Keep weekly review cadence", "Operations Lead", "2026-05-24", "Improves visibility", "Send updated invite"],
            ],
        },
    },
    "client-kickoff-meeting-pack": {
        "title": "Client Kickoff Meeting Pack",
        "audience": "Consultants, agencies, client success teams, implementation leads, and account managers.",
        "summary": "A polished onboarding and kickoff toolkit for starting client work with clearer expectations, smoother facilitation, and stronger follow-through.",
        "value_props": [
            "Client-facing scripts and talk tracks that sound calm and professional",
            "Editable kickoff documents plus working planning files",
            "Example and blank trackers for stakeholder and first-30-days setup",
        ],
        "use_case": "Use this pack when a new client engagement needs a stronger first impression, cleaner operating rhythm, and more visible next steps.",
        "preview_images": ["client-kickoff-workbook.png", "kickoff-deck.png"],
        "csv_defs": {
            "roles-and-responsibilities-matrix": [
                ["Workstream", "Responsible", "Accountable", "Consulted", "Informed"],
                ["Project planning", "Delivery Lead", "Account Lead", "Client Sponsor", "Core Team"],
                ["Decision approvals", "Account Lead", "Client Sponsor", "Delivery Lead", "Project Team"],
            ],
            "stakeholder-register": [
                ["Name", "Role", "Team", "Decision Influence", "Update Need", "Notes"],
                ["Jordan Smith", "Executive Sponsor", "Client Leadership", "High", "Weekly summary", "Primary approval owner"],
                ["Taylor Lee", "Operations Lead", "Client Ops", "Medium", "Action-focused notes", "Needs milestone visibility"],
            ],
            "first-30-days-plan": [
                ["Week", "Focus", "Owner", "Success Signal", "Notes"],
                ["Week 1", "Kickoff alignment and role clarity", "Delivery Lead", "Owners and next steps confirmed", ""],
                ["Week 2", "Discovery inputs and decision path", "Account Lead", "Dependencies surfaced", ""],
            ],
            "decision-tracker": [
                ["Decision", "Owner", "Date Needed", "Impact", "Status"],
                ["Confirm approval path", "Client Sponsor", "2026-05-31", "Affects timeline confidence", "Open"],
            ],
        },
    },
    "website-redesign-planning-kit": {
        "title": "Website Redesign Planning Kit",
        "audience": "Small business owners, consultants, marketing agencies, website agencies, and founders.",
        "summary": "A premium website planning toolkit for clarifying messaging, organizing content, and making redesign or refresh work easier to manage.",
        "value_props": [
            "Useful for selling or planning website work with small businesses",
            "Editable planning files for messaging, content, structure, and launch readiness",
            "Example trackers and worksheets that help teams get clearer faster",
        ],
        "use_case": "Use this kit before a redesign, website refresh, proposal, or content planning sprint when the website needs to say the right thing more clearly.",
        "preview_images": ["website-planning-workbook.png", "website-deck.png"],
        "csv_defs": {
            "website-content-inventory": [
                ["Page", "Current Status", "Keep / Rewrite / Remove", "Primary CTA", "Proof Needed", "Notes"],
                ["Homepage", "Outdated", "Rewrite", "Book a call", "Service proof", "Headline feels vague"],
                ["Services", "Mixed", "Rewrite", "Request review", "Examples and outcomes", ""],
            ],
            "competitor-review-worksheet": [
                ["Competitor", "Main Message", "Offer Clarity", "Proof Quality", "CTA Quality", "Notes"],
                ["Competitor A", "Strong", "Clear", "High", "Medium", "Simple service framing"],
                ["Competitor B", "Weak", "Generic", "Low", "Low", "No strong CTA path"],
            ],
            "website-project-plan": [
                ["Workstream", "Owner", "Target Date", "Status", "Notes"],
                ["Messaging brief", "Strategy Lead", "2026-06-03", "Open", ""],
                ["Content collection", "Client Owner", "2026-06-10", "Open", ""],
            ],
            "launch-readiness-checklist": [
                ["Checklist Item", "Owner", "Status", "Notes"],
                ["Homepage approved", "Client Owner", "Open", ""],
                ["Forms tested", "Developer", "Open", ""],
            ],
        },
    },
    "consulting-proposal-starter-kit": {
        "title": "Consulting Proposal Starter Kit",
        "audience": "Consultants, boutique agencies, fractional operators, and professional services firms.",
        "summary": "A proposal and conversion toolkit for turning discovery calls into cleaner scopes, stronger pricing conversations, and more confident next steps.",
        "value_props": [
            "Professional proposal language without sounding bloated or overly corporate",
            "Editable proposal, onboarding, and follow-up assets",
            "Example and blank commercial trackers for pricing and scope control",
        ],
        "use_case": "Use this kit when you need to move from exploratory conversations into a clearer scope, stronger proposal, and smoother client conversion path.",
        "preview_images": ["consulting-pricing-and-scope-workbook.png", "consulting-deck.png"],
        "csv_defs": {
            "pricing-options-template": [
                ["Option", "Description", "Fee", "Billing Model", "Notes"],
                ["Starter", "Targeted advisory sprint", "$2,500", "Fixed fee", "Good for small planning engagements"],
                ["Core", "Structured delivery support", "$5,500", "Fixed fee", "Adds communication and governance support"],
            ],
            "client-fit-scorecard": [
                ["Criteria", "Score", "Notes"],
                ["Decision access", "4", "Strong sponsor availability"],
                ["Urgency", "5", "Decision needed within two weeks"],
            ],
            "scope-assumptions-tracker": [
                ["Area", "Assumption", "Risk if false", "Owner", "Notes"],
                ["Inputs", "Client content will be available on time", "Timeline shift", "Client Owner", ""],
                ["Approvals", "One decision owner is confirmed", "Rework and delay", "Account Lead", ""],
            ],
        },
    },
    "full-clarpoint-business-execution-toolkit": {
        "title": "Full Clarpoint Business Execution Toolkit",
        "audience": "Consultants, agencies, operators, project managers, founders, and small business owners who want a reusable operating system.",
        "summary": "The full Clarpoint delivery system bundle that combines project structure, kickoff, website planning, proposal support, and executive communication in one flagship package.",
        "value_props": [
            "A full client delivery system instead of isolated templates",
            "Organized downloadable assets by format plus bundled ZIPs",
            "Clear guidance on when to use each toolkit and how they work together",
        ],
        "use_case": "Use this bundle when you want one repeatable system for winning work, starting work, running work, and communicating progress more professionally.",
        "preview_images": ["clarpoint-toolkit-index.png", "bundle-deck.png"],
        "csv_defs": {
            "full-system-file-index": [
                ["Toolkit", "Primary Use", "Formats Included", "Best Time To Use"],
                ["Executive Project Status Pack", "Leadership reporting and delivery visibility", "DOCX, XLSX, PPTX, PDF, CSV", "When active work needs clearer executive updates"],
                ["Client Kickoff Meeting Pack", "Client onboarding and kickoff structure", "DOCX, XLSX, PPTX, PDF, CSV", "Before or right after a new engagement begins"],
            ],
            "toolkit-selection-matrix": [
                ["Scenario", "Best Toolkit", "Why"],
                ["Need cleaner status reporting", "Executive Project Status Pack", "Gives leaders a clearer project readout and control layer"],
                ["Need a stronger kickoff", "Client Kickoff Meeting Pack", "Sets expectations, roles, and next steps more cleanly"],
            ],
            "delivery-workflow-planner": [
                ["Stage", "Toolkit", "Owner", "Notes"],
                ["Proposal", "Consulting Proposal Starter Kit", "Consultant", "Use before scope is finalized"],
                ["Kickoff", "Client Kickoff Meeting Pack", "Delivery Lead", "Use to establish cadence and expectations"],
            ],
        },
    },
}


def ensure_docx(path: Path, title: str, subtitle: str, sections: list[tuple[str, list[str]]]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(15)

    title_p = document.add_paragraph(style="Title")
    title_p.add_run(title).bold = True

    subtitle_p = document.add_paragraph()
    subtitle_p.add_run(subtitle)
    subtitle_p.runs[0].italic = True

    tag_p = document.add_paragraph()
    tag_p.add_run(BRAND["tagline"]).bold = True

    for heading, bullets in sections:
        document.add_paragraph(heading, style="Heading 1")
        for bullet in bullets:
            p = document.add_paragraph(style="List Bullet")
            p.add_run(bullet)

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(BRAND["site"]).italic = True
    document.save(path)


def ensure_pdf(path: Path, title: str, subtitle: str, sections: list[tuple[str, list[str]]]) -> None:
    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        leftMargin=0.72 * inch,
        rightMargin=0.72 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.72 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=22, textColor=colors.HexColor("#162424"), spaceAfter=8)
    subtitle_style = ParagraphStyle("Sub", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, textColor=colors.HexColor("#556766"), leading=14, spaceAfter=8)
    heading_style = ParagraphStyle("Head", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, textColor=colors.HexColor("#0F766E"), spaceBefore=8, spaceAfter=6)
    body_style = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, textColor=colors.HexColor("#223434"), leading=14)
    small_style = ParagraphStyle("Small", parent=styles["BodyText"], fontName="Helvetica", fontSize=9, textColor=colors.HexColor("#5A6A69"), leading=12)

    story = [
        Paragraph(title, title_style),
        Paragraph(subtitle, subtitle_style),
        Paragraph(BRAND["tagline"], small_style),
        Spacer(1, 0.14 * inch),
    ]

    for heading, bullets in sections:
        story.append(Paragraph(heading, heading_style))
        story.append(ListFlowable([ListItem(Paragraph(item, body_style), leftIndent=8) for item in bullets], bulletType="bullet", leftIndent=14))

    story.extend([Spacer(1, 0.16 * inch), Paragraph(BRAND["site"], small_style)])
    doc.build(story)


def write_csv(path: Path, rows: list[list[str]]) -> None:
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerows(rows)


def blank_rows_from_rows(rows: list[list[str]]) -> list[list[str]]:
    headers = rows[0]
    blank = [headers]
    for _ in rows[1:]:
        blank.append(["" for _ in headers])
    return blank


def copy_to_format_dirs(product_root: Path) -> None:
    for file_path in [p for p in product_root.iterdir() if p.is_file()]:
        target_dir = FORMAT_DIRS.get(file_path.suffix.lower())
        if not target_dir:
            continue
        destination = product_root / target_dir / file_path.name
        destination.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(file_path, destination)


def copy_preview_assets(product_root: Path, product: dict) -> list[str]:
    copied = []
    preview_dir = product_root / "preview"
    preview_dir.mkdir(parents=True, exist_ok=True)
    for image_name in product.get("preview_images", []):
        candidates = [
            PREVIEW_OUTPUT / "workbook-previews" / image_name,
            PREVIEW_OUTPUT / "deck-previews" / image_name,
        ]
        for source in candidates:
            if source.exists():
                shutil.copy2(source, preview_dir / image_name)
                copied.append(image_name)
                break
    return copied


def create_preview_html(product_root: Path, product: dict, images: list[str], manifest_paths: list[dict]) -> None:
    preview_dir = product_root / "preview"
    image_markup = "".join(
        f'<figure class="preview-shot"><img src="{name}" alt="{product["title"]} preview image"></figure>'
        for name in images
    )
    file_markup = "".join(
        f'<li><strong>{entry["label"]}</strong><span>{entry["description"]}</span></li>'
        for entry in manifest_paths[:8]
    )
    html = f"""<!doctype html>
<html lang="en">
  <head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>{product['title']} Preview</title>
    <style>
      body {{
        margin: 0;
        font-family: "Manrope", "Aptos", sans-serif;
        background: #f6f1e7;
        color: #162424;
      }}
      .wrap {{
        max-width: 1120px;
        margin: 0 auto;
        padding: 48px 28px 60px;
      }}
      .eyebrow {{
        color: #0f766e;
        font-size: 13px;
        font-weight: 800;
        letter-spacing: 0.04em;
        text-transform: uppercase;
      }}
      h1 {{
        font-size: 46px;
        line-height: 1.05;
        margin: 12px 0 16px;
      }}
      p {{
        max-width: 760px;
        color: #556766;
        font-size: 18px;
        line-height: 1.6;
      }}
      .grid {{
        display: grid;
        grid-template-columns: repeat(2, minmax(0, 1fr));
        gap: 22px;
        margin-top: 28px;
      }}
      .card {{
        background: rgba(255,255,255,0.92);
        border: 1px solid rgba(22,36,36,0.08);
        border-radius: 24px;
        padding: 24px;
        box-shadow: 0 22px 48px rgba(20,34,32,0.09);
      }}
      img {{
        width: 100%;
        display: block;
        border-radius: 16px;
      }}
      ul {{
        margin: 18px 0 0;
        padding: 0;
        list-style: none;
        display: grid;
        gap: 14px;
      }}
      li {{
        display: grid;
        gap: 4px;
        padding: 14px 0;
        border-top: 1px solid rgba(22,36,36,0.08);
      }}
      li:first-child {{
        border-top: 0;
        padding-top: 0;
      }}
      li span {{
        color: #556766;
        font-size: 15px;
        line-height: 1.5;
      }}
      @media (max-width: 860px) {{
        .grid {{
          grid-template-columns: 1fr;
        }}
        h1 {{
          font-size: 36px;
        }}
      }}
    </style>
  </head>
  <body>
    <div class="wrap">
      <p class="eyebrow">Clarpoint Preview</p>
      <h1>{product['title']}</h1>
      <p>{product['summary']}</p>
      <div class="grid">
        <section class="card">
          <p class="eyebrow">What buyers receive</p>
          <ul>{file_markup}</ul>
        </section>
        <section class="card">
          <p class="eyebrow">Preview</p>
          {image_markup or "<p>Workbook and deck previews are included after generation.</p>"}
        </section>
      </div>
    </div>
  </body>
</html>
"""
    (preview_dir / "index.html").write_text(html, encoding="utf-8")


def describe_file(filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return "Polished PDF guide or overview designed for quick review and client-ready sharing."
    if lower.endswith(".docx"):
        return "Editable Word document built for direct reuse, customization, and same-day delivery work."
    if lower.endswith(".xlsx"):
        return "Formatted Excel workbook with structured tabs, sample data, and reusable working views."
    if lower.endswith(".pptx"):
        return "Presentation deck for client, leadership, kickoff, or delivery conversations."
    if lower.endswith(".csv"):
        return "Import-friendly CSV file with either realistic sample rows or a blank version ready to customize."
    if lower.endswith(".html"):
        return "Optional browser preview that shows the toolkit structure and included files."
    return "Supporting file included in the product download."


def build_manifest(product_root: Path, product: dict) -> list[dict]:
    files = []
    for directory in ["pdf", "word", "excel", "powerpoint", "csv", "preview", "docs"]:
        for file_path in sorted((product_root / directory).glob("*")):
            if not file_path.is_file():
                continue
            relative_path = file_path.relative_to(product_root).as_posix()
            files.append(
                {
                    "format": directory,
                    "path": relative_path,
                    "label": file_path.name,
                    "description": describe_file(file_path.name),
                }
            )

    manifest = {
        "productId": product_root.name,
        "productName": product["title"],
        "price": None,
        "files": files,
    }
    (product_root / "downloadManifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return files


def build_zip(product_root: Path) -> None:
    zip_path = product_root / "download-all.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for directory in ["pdf", "word", "excel", "powerpoint", "csv", "preview", "docs"]:
            for file_path in sorted((product_root / directory).glob("*")):
                if file_path.is_file():
                    archive.write(file_path, arcname=file_path.relative_to(product_root))
        manifest = product_root / "downloadManifest.json"
        if manifest.exists():
            archive.write(manifest, arcname="downloadManifest.json")


def create_support_docs(product_root: Path, slug: str, product: dict) -> None:
    docs_dir = product_root / "docs"
    docs_dir.mkdir(parents=True, exist_ok=True)

    usage_sections = [
        ("Who this toolkit is for", [product["audience"]]),
        ("Best use cases", [product["use_case"]] + product["value_props"]),
        ("Recommended working sequence", [
            "Start with the PDF overview and usage guide to understand the structure of the pack.",
            "Use the editable Word files for communication, scripts, and client-facing language.",
            "Use the Excel workbook and CSV files for tracking, importing, or sharing structured data.",
            "Use the PowerPoint deck when you need a cleaner readout for clients, sponsors, or leadership groups.",
        ]),
    ]
    ensure_docx(docs_dir / "Usage Guide.docx", f"{product['title']} | Usage Guide", product["summary"], usage_sections)
    ensure_pdf(product_root / "pdf" / "Quick Start Guide.pdf", f"{product['title']} | Quick Start Guide", product["summary"], usage_sections)

    example_sections = [
        ("Sample scenario", [
            f"This example shows how {product['title']} can be used to organize work, improve communication, and reduce follow-up noise in a real delivery situation.",
            "Review the example first if you want to see how the files can work together before replacing the sample content with your own.",
        ]),
        ("Suggested first moves", [
            "Replace sample names, dates, and owners with your real project or client information.",
            "Keep the structure intact at first so the toolkit stays easy to review and reuse.",
        ]),
    ]
    ensure_docx(docs_dir / "Example Use Case.docx", f"{product['title']} | Example Use Case", product["use_case"], example_sections)

    license_sections = [
        ("License summary", [
            "This purchase is licensed for use inside your own business, team, or client service delivery work.",
            "You may customize the files for your company or your clients, but you may not resell, redistribute, or repackage the toolkit as your own digital product.",
            f"For questions about broader licensing or customization support, contact {BRAND['site']}.",
        ]),
    ]
    ensure_docx(docs_dir / "Licensing Note.docx", f"{product['title']} | Licensing Note", "Simple commercial use guidance for buyers.", license_sections)


def build_csv_variants(product_root: Path, product: dict) -> None:
    csv_dir = product_root / "csv"
    csv_dir.mkdir(parents=True, exist_ok=True)

    for base_name, rows in product.get("csv_defs", {}).items():
        write_csv(csv_dir / f"{base_name}-example.csv", rows)
        write_csv(csv_dir / f"{base_name}-blank.csv", blank_rows_from_rows(rows))

    for top_level_csv in product_root.glob("*.csv"):
        target = csv_dir / top_level_csv.name
        shutil.copy2(top_level_csv, target)
        rows = list(csv.reader(top_level_csv.open("r", encoding="utf-8")))
        stem = top_level_csv.stem
        example_path = csv_dir / f"{stem}-example.csv"
        blank_path = csv_dir / f"{stem}-blank.csv"
        if not example_path.exists():
            write_csv(example_path, rows)
        if not blank_path.exists():
            write_csv(blank_path, blank_rows_from_rows(rows))


def main() -> None:
    for slug, product in PRODUCTS.items():
        product_root = DOWNLOADS / slug
        if not product_root.exists():
            continue

        for directory in ["pdf", "word", "excel", "powerpoint", "csv", "preview", "docs"]:
            shutil.rmtree(product_root / directory, ignore_errors=True)
            (product_root / directory).mkdir(parents=True, exist_ok=True)

        copy_to_format_dirs(product_root)
        build_csv_variants(product_root, product)
        create_support_docs(product_root, slug, product)
        images = copy_preview_assets(product_root, product)
        manifest_files = build_manifest(product_root, product)
        create_preview_html(product_root, product, images, manifest_files)
        manifest_files = build_manifest(product_root, product)
        build_zip(product_root)


if __name__ == "__main__":
    main()
