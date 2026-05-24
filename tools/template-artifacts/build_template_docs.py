from __future__ import annotations

import csv
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[2]
DOWNLOADS = ROOT / "public" / "downloads" / "templates"

BRAND = {
    "name": "Clarpoint",
    "tagline": "Clear plans. Stronger communication. Better execution.",
}


def clean(text: str) -> str:
    return textwrap.dedent(text).strip() + "\n"


PACKS = {
    "executive-project-status-pack": {
        "title": "Executive Project Status Pack",
        "guide_filename": "clarpoint-executive-status-pack-guide.docx",
        "description": "A leadership-ready reporting toolkit for project managers, consultants, and delivery leads who need shorter updates, clearer decisions, and stronger delivery visibility.",
        "audience": "Project managers, PMO leads, consultants, delivery managers, program leads",
        "outcomes": [
            "Give leaders a fast, accurate read on project status without over-explaining.",
            "Keep risks, actions, decisions, and escalations visible in one operating set.",
            "Reduce the time it takes to prepare weekly readouts and steering committee updates.",
        ],
        "workflow": [
            "Start with the Executive Status Report Template to frame the reporting story.",
            "Use the workbook and CSV logs to keep risks, actions, and decisions current.",
            "Use the email and escalation templates when leadership or client communication needs to move quickly.",
            "Use the steering committee deck to turn active project work into an executive discussion.",
        ],
        "deliverables": [
            ("executive-project-control-workbook.xlsx", "Multi-sheet workbook for status tracking, RAID management, milestones, and decision visibility."),
            ("executive-status-report-template.docx", "Editable one-page project status report for weekly or biweekly leadership reporting."),
            ("executive-update-email-template.docx", "Short-form leadership update and decision request email template."),
            ("escalation-communication-template.docx", "Structured escalation note for delivery risk, blockers, ownership gaps, or timeline pressure."),
            ("steering-committee-summary-deck.pptx", "Presentation deck for steering committee or executive checkpoint discussions."),
            ("raid-log.csv", "Lightweight tracker for risks, assumptions, issues, and dependencies."),
            ("action-item-tracker.csv", "Owner-based follow-up tracker with due dates and status."),
            ("decision-log.csv", "Decision history tracker with owners, impact, and follow-up needs."),
            ("steering-committee-slide-outline.md", "Narrative outline behind the editable presentation deck."),
        ],
        "docs": {
            "README.md": """
            # Executive Project Status Pack

            This pack gives you a practical reporting system for active project work. It is designed to help you move from scattered status notes to clear weekly reporting, cleaner escalation, and more useful executive conversations.

            ## What this pack helps you do
            - Frame project status in a way leaders can absorb quickly.
            - Keep risks, actions, and decisions in a usable control layer.
            - Escalate issues with a clearer statement of impact, owner, and next step.
            - Walk into steering committee reviews with a tighter decision narrative.
            """,
            "executive-status-report-template.md": """
            # Executive Status Report Template

            ## Reporting Snapshot
            - Project / workstream:
            - Reporting period:
            - Executive sponsor:
            - Delivery lead:
            - Overall status (Green / Yellow / Red):
            - Confidence level:

            ## What changed this period
            - Key progress made:
            - Milestones completed:
            - Notable shifts in scope, dependency, cost, or timeline:

            ## What leadership should know now
            - Top message:
            - Why it matters:
            - What happens next if nothing changes:

            ## Current priorities
            - Priority 1:
            - Priority 2:
            - Priority 3:

            ## Milestones and timing
            - Upcoming milestone:
            - Target date:
            - Risk to date, if any:

            ## Risks, blockers, and decisions
            - Top risk:
            - Blocker or dependency:
            - Decision needed:
            - Decision owner:

            ## Next seven days
            - Action 1:
            - Action 2:
            - Action 3:

            ## Support requested
            - What support is needed:
            - From whom:
            - By when:
            """,
            "executive-update-email-template.md": """
            # Executive Update Email Template

            ## Subject line
            [Project / Initiative] weekly update | status, risks, and decisions

            ## Email body
            Hi [Name / team],

            Here is the current view of [project / initiative] for the week of [date].

            ### Overall status
            [Green / Yellow / Red] because [short reason].

            ### What moved forward
            - [Progress item]
            - [Progress item]
            - [Progress item]

            ### What needs attention
            - [Risk or blocker]
            - [Dependency or delayed decision]

            ### Decision or support needed
            - Ask:
            - Owner:
            - Needed by:

            ### Next step
            [One clear next move for the team or leadership group.]

            Thanks,
            [Name]
            """,
            "escalation-communication-template.md": """
            # Escalation Communication Template

            ## Escalation summary
            - Issue title:
            - Date raised:
            - Owner:
            - Affected workstream:
            - Severity:

            ## What is happening
            [Write a short description of the issue, including the current state and why it is no longer manageable through normal working channels.]

            ## Business impact
            - Timeline impact:
            - Delivery impact:
            - Client or stakeholder impact:
            - Risk if unresolved:

            ## What has already been tried
            - Action taken:
            - Outcome:
            - Why it did not fully resolve the issue:

            ## What is needed now
            - Decision, support, or owner intervention required:
            - Required by:
            - Recommended path:

            ## Immediate next step
            [One specific action that should happen next.]
            """,
            "steering-committee-slide-outline.md": """
            # Steering Committee Summary Deck Outline

            ## Slide 1: Current view
            - Initiative name
            - Overall status
            - One-sentence leadership message

            ## Slide 2: Progress since last review
            - Milestones completed
            - Work in motion
            - Material changes

            ## Slide 3: Risks and decisions
            - Top risks
            - Key dependencies
            - Decisions needed

            ## Slide 4: Timeline and next steps
            - Current milestone path
            - What happens in the next reporting period
            - Support requested from sponsors or steering committee members
            """,
        },
        "csvs": {
            "raid-log.csv": [
                ["ID", "Type", "Description", "Impact", "Owner", "Mitigation / Next Step", "Target Date", "Status"],
                ["R-01", "Risk", "Approval timing may impact build start", "Medium", "Program Manager", "Confirm approval owner and decision date", "2026-06-02", "Open"],
                ["A-01", "Assumption", "Content inputs will be available on time", "Medium", "Business Lead", "Validate content readiness in weekly call", "2026-06-04", "Open"],
                ["I-01", "Issue", "Vendor dependency has not confirmed delivery date", "High", "Delivery Lead", "Escalate in steering prep note", "2026-05-30", "In Progress"],
                ["D-01", "Dependency", "Analytics tagging requires external support", "Medium", "Tech Lead", "Book working session with analytics owner", "2026-06-06", "Open"],
            ],
            "action-item-tracker.csv": [
                ["Action ID", "Action", "Owner", "Due Date", "Status", "Priority", "Notes"],
                ["ACT-01", "Confirm milestone owners for next release", "Program Manager", "2026-05-29", "In Progress", "High", "Review in Monday update"],
                ["ACT-02", "Draft decision note for steering committee", "Delivery Lead", "2026-05-31", "Open", "High", "Needs sponsor context"],
                ["ACT-03", "Validate client-facing launch risks", "Account Lead", "2026-06-03", "Open", "Medium", "Coordinate with tech lead"],
            ],
            "decision-log.csv": [
                ["Decision ID", "Decision", "Date", "Decision Owner", "Reason", "Impact", "Follow-Up"],
                ["DEC-01", "Keep phased launch approach", "2026-05-24", "Executive Sponsor", "Reduces near-term implementation risk", "Scope / timeline", "Update milestone plan"],
                ["DEC-02", "Escalate content dependency", "2026-05-24", "Program Manager", "Content readiness threatens launch date", "Stakeholder communication", "Raise in leadership update"],
            ],
        },
    },
    "raid-log-action-tracker-bundle": {
        "title": "RAID Log + Action Tracker Bundle",
        "guide_filename": "clarpoint-raid-action-bundle-guide.docx",
        "description": "A simple control bundle that gives teams one place to track risks, issues, decisions, owners, due dates, and follow-up commitments.",
        "audience": "Project leads, operators, account leads, delivery managers, small internal teams",
        "outcomes": [
            "Give the team a visible source of truth for project control items.",
            "Make owner-based follow-up easier to review in weekly meetings.",
            "Reduce the number of decisions and actions buried in inboxes or meeting notes.",
        ],
        "workflow": [
            "Use the workbook as the master version for active work.",
            "Share the CSV versions if the team wants lighter files for import or ad hoc updates.",
            "Use the weekly follow-up template to run shorter, more accountable check-ins.",
        ],
        "deliverables": [
            ("raid-action-control-workbook.xlsx", "Editable workbook with RAID, actions, decisions, and a meeting-ready follow-up view."),
            ("raid-log.csv", "Simple tracker for risks, assumptions, issues, and dependencies."),
            ("action-tracker.csv", "Action register with owners, due dates, and status."),
            ("decision-log.csv", "Decision log with rationale and downstream follow-up."),
            ("weekly-follow-up-template.docx", "Weekly review template for moving open items toward action."),
        ],
        "docs": {
            "README.md": """
            # RAID Log + Action Tracker Bundle

            This bundle gives you a practical control layer for active work. Use it when there is motion, but the team does not have a clean place to track commitments, decisions, and delivery risks.
            """,
            "weekly-follow-up-template.md": """
            # Weekly Follow-Up Template

            ## Meeting purpose
            Use this template to move open actions, risks, and decisions forward with less noise.

            ## Current view
            - Reporting week:
            - Facilitator:
            - Team:

            ## Open actions to review
            - Action:
            - Owner:
            - Due date:
            - What is blocking progress:

            ## Risks or issues that need attention
            - Risk / issue:
            - Impact:
            - Mitigation:
            - Escalation needed:

            ## Decisions required
            - Decision:
            - Owner:
            - Deadline:
            - If delayed, impact is:

            ## Commitments from this session
            - New action:
            - Owner:
            - Due date:
            """,
        },
        "csvs": {
            "raid-log.csv": [
                ["ID", "Category", "Item", "Impact", "Owner", "Mitigation", "Review Date", "Status"],
                ["R-01", "Risk", "Requirements still shifting", "High", "Project Lead", "Freeze core scope by next checkpoint", "2026-05-29", "Open"],
                ["I-01", "Issue", "Decision backlog slowing approvals", "High", "Sponsor", "Prioritize open approvals in steering prep", "2026-05-30", "In Progress"],
            ],
            "action-tracker.csv": [
                ["Action ID", "Action", "Owner", "Due Date", "Status", "Notes"],
                ["ACT-01", "Close outstanding scope questions", "Project Lead", "2026-05-28", "Open", "Needs client input"],
                ["ACT-02", "Update weekly risk summary", "PM", "2026-05-29", "In Progress", "Prepare for Friday review"],
            ],
            "decision-log.csv": [
                ["Decision ID", "Decision", "Decision Owner", "Date", "Impact", "Next Step"],
                ["DEC-01", "Keep weekly review cadence", "Operations Lead", "2026-05-24", "Improves visibility", "Send updated invite"],
            ],
        },
    },
    "client-kickoff-meeting-pack": {
        "title": "Client Kickoff Meeting Pack",
        "guide_filename": "clarpoint-client-kickoff-guide.docx",
        "description": "A professional kickoff toolkit for consultants, agencies, and delivery teams who want cleaner starts, clearer expectations, and stronger follow-through.",
        "audience": "Consultants, agencies, account leads, client success teams, implementation teams",
        "outcomes": [
            "Start new client work with a clearer scope and better shared expectations.",
            "Create a usable record of stakeholder input, roles, and next steps.",
            "Make the first client touchpoint feel structured, calm, and professional.",
        ],
        "workflow": [
            "Use the agenda and deck to shape the kickoff session.",
            "Capture stakeholder inputs and role expectations in the intake and RACI documents.",
            "Use the meeting notes, workbook, and follow-up email to lock in commitments after the call.",
        ],
        "deliverables": [
            ("client-kickoff-deck.pptx", "Editable kickoff deck for the meeting itself."),
            ("client-kickoff-workbook.xlsx", "Stakeholder and action workbook for kickoff follow-up."),
            ("kickoff-agenda.docx", "Structured agenda for the kickoff session."),
            ("stakeholder-intake-form.docx", "Question set for collecting stakeholder goals, constraints, and expectations."),
            ("roles-and-responsibilities-template.docx", "Simple ownership template for delivery roles and working relationships."),
            ("meeting-notes-template.docx", "Reusable kickoff notes and action capture template."),
            ("follow-up-email-template.docx", "Post-kickoff email template that confirms decisions and next steps."),
        ],
        "docs": {
            "README.md": """
            # Client Kickoff Meeting Pack

            This pack helps you start new client work with clearer structure, better stakeholder alignment, and a stronger first impression. It is designed for teams that want a more polished kickoff without overcomplicating the process.
            """,
            "kickoff-agenda.md": """
            # Client Kickoff Agenda

            ## 1. Welcome and introductions
            - Introduce core team members
            - Confirm sponsor, day-to-day leads, and working team

            ## 2. Why this work matters
            - Desired outcome
            - Business context
            - What success should look like

            ## 3. Scope and priorities
            - In scope
            - Out of scope
            - Known assumptions

            ## 4. Roles and responsibilities
            - Client owner
            - Clarpoint / delivery owner
            - Decision makers
            - Reviewers / contributors

            ## 5. Ways of working
            - Meeting cadence
            - Communication channels
            - Approvals and escalation path

            ## 6. Risks, dependencies, and questions
            - Open questions
            - Constraints
            - Immediate dependencies

            ## 7. Next steps
            - Actions
            - Owners
            - Due dates
            """,
            "kickoff-deck-outline.md": """
            # Client Kickoff Deck Outline

            ## Slide 1: Engagement overview
            - Project name
            - Outcome statement
            - Working team

            ## Slide 2: Scope and success
            - What is in scope
            - What success looks like
            - What is out of scope

            ## Slide 3: Roles and ways of working
            - Client owner
            - Delivery owner
            - Cadence
            - Decision path

            ## Slide 4: Next steps
            - Immediate actions
            - Owners
            - Target dates
            """,
            "stakeholder-intake-form.md": """
            # Stakeholder Intake Form

            ## Stakeholder details
            - Name:
            - Role:
            - Team:
            - Contact preference:

            ## Business context
            - What is most important about this work?
            - What pressure or risk is driving urgency?
            - What needs to be true for this to feel successful?

            ## Operating needs
            - What decisions should involve you?
            - What updates do you expect to receive?
            - What are your main concerns or constraints?

            ## Delivery considerations
            - Key milestone or date:
            - Known dependencies:
            - Preferred way of working:
            """,
            "roles-and-responsibilities-template.md": """
            # Roles and Responsibilities Template

            ## Core roles
            - Executive sponsor:
            - Business owner:
            - Delivery lead:
            - Project manager:
            - Technical owner:
            - Client approver:

            ## Responsibility model
            - Responsible:
            - Accountable:
            - Consulted:
            - Informed:

            ## Review points
            - Weekly action review:
            - Decision review:
            - Escalation path:
            """,
            "meeting-notes-template.md": """
            # Kickoff Meeting Notes Template

            ## Meeting details
            - Date:
            - Facilitator:
            - Attendees:

            ## Main discussion points
            - Objective:
            - Scope:
            - Key context:

            ## Decisions made
            - Decision:
            - Owner:

            ## Actions
            - Action:
            - Owner:
            - Due date:

            ## Follow-up needed
            - Open question:
            - Next move:
            """,
            "follow-up-email-template.md": """
            # Kickoff Follow-Up Email Template

            ## Subject line
            [Project / engagement] kickoff recap and next steps

            ## Email body
            Hi [Name / team],

            Thanks again for today’s kickoff. Below is the current summary of what we aligned on.

            ### Agreed objective
            [Insert short outcome statement.]

            ### Main points confirmed
            - [Point]
            - [Point]
            - [Point]

            ### Actions and owners
            - [Action] | [Owner] | [Due date]
            - [Action] | [Owner] | [Due date]

            ### Open questions
            - [Question]

            ### Next checkpoint
            [Insert date / meeting / review point.]

            Thanks,
            [Name]
            """,
        },
        "csvs": {},
    },
    "website-redesign-planning-kit": {
        "title": "Website Redesign Planning Kit",
        "guide_filename": "clarpoint-website-planning-guide.docx",
        "description": "A practical planning toolkit for businesses preparing for a website redesign, content refresh, or cleaner digital presence.",
        "audience": "Small business owners, founders, consultants, marketing leads, local businesses",
        "outcomes": [
            "Clarify what the website needs to say and how it should support the business.",
            "Reduce guesswork around page planning, content needs, and launch readiness.",
            "Create a more credible digital presence with clearer messaging and stronger calls to action.",
        ],
        "workflow": [
            "Use the questionnaire and messaging worksheet first to clarify business positioning.",
            "Use the content planner and competitor worksheet to shape the site structure.",
            "Use the launch and maintenance checklists to turn planning into a cleaner rollout.",
        ],
        "deliverables": [
            ("website-planning-workbook.xlsx", "Planning workbook with page structure, content needs, and launch tracking."),
            ("website-planning-brief-deck.pptx", "Editable presentation deck for website planning, message alignment, and launch framing."),
            ("website-intake-questionnaire.docx", "Business and website discovery prompts."),
            ("website-content-planner.docx", "Page-by-page content planning template."),
            ("homepage-messaging-worksheet.docx", "Homepage headline, proof, CTA, and offer clarity worksheet."),
            ("competitor-review-worksheet.docx", "Lightweight competitor positioning review template."),
            ("website-launch-checklist.docx", "Pre-launch and go-live checklist."),
            ("website-maintenance-checklist.docx", "Ongoing content, SEO, and operations checklist."),
        ],
        "docs": {
            "README.md": """
            # Website Redesign Planning Kit

            This kit helps businesses get clearer before a redesign or website refresh. It gives you a practical planning structure for messaging, content, launch preparation, and maintenance.
            """,
            "website-intake-questionnaire.md": """
            # Website Intake Questionnaire

            ## Business basics
            - Business name:
            - Primary service or offer:
            - Main audience:
            - Geographic market, if relevant:

            ## Why the website needs attention
            - What feels unclear or outdated today?
            - What do visitors struggle to understand?
            - What should the website help people do?

            ## Offer and positioning
            - What are your core services?
            - What makes your business credible?
            - What should a visitor remember after 30 seconds?

            ## Conversion goals
            - Primary CTA:
            - Secondary CTA:
            - Best lead type:

            ## Project constraints
            - Timing:
            - Dependencies:
            - Existing assets or brand rules:
            """,
            "website-content-planner.md": """
            # Website Content Planner

            ## Page planning table
            - Page name:
            - Main purpose:
            - Key message:
            - Supporting proof:
            - Primary CTA:
            - Assets needed:

            ## Content review prompts
            - What should this page explain first?
            - What should the visitor do next?
            - What proof should be visible here?
            """,
            "homepage-messaging-worksheet.md": """
            # Homepage Messaging Worksheet

            ## Headline
            - What does the business help with?
            - Who is it for?
            - What outcome should be clear in one line?

            ## Supporting copy
            - What problem does the visitor have right now?
            - Why is your business credible to solve it?
            - What makes the offer feel clear and low-friction?

            ## Proof and trust
            - Experience indicators:
            - Process or approach:
            - Testimonials or signals:

            ## CTA structure
            - Primary CTA:
            - Secondary CTA:
            - Why would someone click now?
            """,
            "competitor-review-worksheet.md": """
            # Competitor Review Worksheet

            ## Competitor details
            - Competitor:
            - Website URL:
            - Main offer:

            ## What they do well
            - Clear headline:
            - Strong proof:
            - Good CTA:

            ## What feels weak or generic
            - Unclear message:
            - Weak offer framing:
            - Missing trust elements:

            ## What Clarpoint / your business should do differently
            - Message angle:
            - Offer clarity:
            - Conversion improvement:
            """,
            "website-launch-checklist.md": """
            # Website Launch Checklist

            ## Before launch
            - [ ] Final copy approved
            - [ ] Contact forms tested
            - [ ] Mobile layout reviewed
            - [ ] CTA links checked
            - [ ] Metadata updated
            - [ ] Analytics / tracking confirmed

            ## Launch day
            - [ ] DNS / hosting confirmed
            - [ ] Final smoke test complete
            - [ ] Core pages reviewed live
            - [ ] Team knows where to report issues

            ## After launch
            - [ ] Capture first round of feedback
            - [ ] Fix priority issues
            - [ ] Confirm lead flow works
            """,
            "website-maintenance-checklist.md": """
            # Website Maintenance Checklist

            ## Monthly review
            - [ ] Check contact forms
            - [ ] Review homepage clarity
            - [ ] Update service or offer changes
            - [ ] Confirm key links still work

            ## Quarterly review
            - [ ] Refresh proof / examples
            - [ ] Review SEO basics and metadata
            - [ ] Check page performance on mobile
            - [ ] Look for outdated language or assets

            ## Ongoing
            - [ ] Add new work examples
            - [ ] Keep CTAs aligned to current offer
            - [ ] Keep ownership for updates clear
            """,
        },
        "csvs": {},
    },
    "consulting-proposal-starter-kit": {
        "title": "Consulting Proposal Starter Kit",
        "guide_filename": "clarpoint-consulting-proposal-guide.docx",
        "description": "A proposal and onboarding toolkit for consultants, freelancers, and boutique agencies who want cleaner client setup, pricing structure, and scope control.",
        "audience": "Consultants, freelancers, boutique agencies, independent operators",
        "outcomes": [
            "Speed up proposal writing without starting from a blank page.",
            "Make onboarding and scope expectations clearer from the start.",
            "Create more consistency across pricing, discovery, and scope management.",
        ],
        "workflow": [
            "Use the service offering worksheet and discovery notes to shape the proposal.",
            "Use the pricing table and workbook to frame fees, scope, and options.",
            "Use the onboarding checklist and scope change request template after the deal is live.",
        ],
        "deliverables": [
            ("consulting-pricing-and-scope-workbook.xlsx", "Editable workbook for pricing, scope structure, and onboarding checkpoints."),
            ("consulting-proposal-template.docx", "Proposal template for service-based work."),
            ("service-offering-worksheet.docx", "Worksheet for clarifying offers, outcomes, and delivery approach."),
            ("pricing-table.csv", "Simple pricing table you can adapt for fixed fee or phased work."),
            ("discovery-call-notes-template.docx", "Discovery call note template for qualifying work and next steps."),
            ("client-onboarding-checklist.docx", "Checklist for new client setup."),
            ("scope-change-request-template.docx", "Structured change request template for scope or timeline shifts."),
        ],
        "docs": {
            "README.md": """
            # Consulting Proposal Starter Kit

            This kit helps you create cleaner proposals and reduce onboarding noise. It is built for consultants and smaller service businesses that want a more polished sales-to-delivery handoff.
            """,
            "consulting-proposal-template.md": """
            # Consulting Proposal Template

            ## Cover
            - Client:
            - Project / engagement:
            - Date:
            - Prepared by:

            ## Executive summary
            [Summarize the client need, the proposed support, and the main outcome.]

            ## Current challenge
            - What is happening now:
            - Why it matters:
            - What is getting slowed down:

            ## Scope of support
            - Workstream 1:
            - Workstream 2:
            - Workstream 3:

            ## Deliverables
            - Deliverable:
            - Deliverable:
            - Deliverable:

            ## Timeline and cadence
            - Start date:
            - Duration:
            - Check-in rhythm:

            ## Investment
            - Fee:
            - Billing approach:
            - Assumptions:

            ## Next steps
            - Approval process:
            - Kickoff timing:
            """,
            "service-offering-worksheet.md": """
            # Service Offering Worksheet

            ## Offer basics
            - Service name:
            - Who it is for:
            - Main business problem solved:
            - Outcome the client should expect:

            ## Scope framing
            - Core deliverables:
            - Out-of-scope items:
            - Dependencies:
            - Success markers:

            ## Delivery model
            - Sprint, retainer, project, or advisory:
            - Meeting cadence:
            - Review points:
            """,
            "discovery-call-notes-template.md": """
            # Discovery Call Notes Template

            ## Client details
            - Client:
            - Contact:
            - Date:

            ## Current need
            - What prompted the conversation?
            - What is not working today?
            - Why now?

            ## Business context
            - Stakeholders:
            - Timeline:
            - Constraints:

            ## Opportunity fit
            - Best support model:
            - Risks to close:
            - Follow-up needed:

            ## Next step
            - Proposal:
            - Follow-up call:
            - Internal action:
            """,
            "client-onboarding-checklist.md": """
            # Client Onboarding Checklist

            ## Before kickoff
            - [ ] Signed agreement received
            - [ ] Primary contacts confirmed
            - [ ] Shared folders / tools created
            - [ ] Kickoff scheduled

            ## Kickoff setup
            - [ ] Agenda sent
            - [ ] Scope summary prepared
            - [ ] Roles confirmed
            - [ ] Communication cadence confirmed

            ## First week
            - [ ] Working plan sent
            - [ ] Owners confirmed
            - [ ] Immediate risks captured
            - [ ] First follow-up note sent
            """,
            "scope-change-request-template.md": """
            # Scope Change Request Template

            ## Change summary
            - Requested by:
            - Date:
            - Change title:

            ## Requested change
            [Describe what is changing and why.]

            ## Impact review
            - Timeline impact:
            - Cost / effort impact:
            - Resourcing impact:
            - Delivery risk impact:

            ## Recommendation
            - Recommended response:
            - Approval owner:
            - Next step:
            """,
        },
        "csvs": {
            "pricing-table.csv": [
                ["Package", "Description", "Fee", "Billing Model", "Notes"],
                ["Starter", "Targeted advisory or sprint support", "$2,500", "Fixed fee", "Good for short planning engagements"],
                ["Core", "Structured delivery and communication support", "$5,500", "Fixed fee", "Add reporting and governance"],
                ["Embedded", "Ongoing weekly support and follow-through", "$3,500 / month", "Monthly retainer", "Best for steady leadership support"],
            ]
        },
    },
    "full-clarpoint-business-execution-toolkit": {
        "title": "Full Clarpoint Business Execution Toolkit",
        "guide_filename": "clarpoint-business-execution-bundle-guide.docx",
        "description": "The complete Clarpoint bundle for project structure, executive communication, client delivery, proposal setup, and website planning.",
        "audience": "Growing businesses, consultants, operators, founders, delivery teams",
        "outcomes": [
            "Give buyers a complete operating starter system instead of one isolated template.",
            "Bundle the most useful project, consulting, and website planning tools in one place.",
            "Make it easier to move from opportunity, to plan, to delivery, to client communication.",
        ],
        "workflow": [
            "Use the included product guides to pick the right pack for the problem you are solving first.",
            "Start with the project, client, or website planning pack that fits your current need.",
            "Use the bonus executive communication templates when leadership visibility or client confidence matters.",
        ],
        "deliverables": [
            ("executive-project-status-pack.zip", "Full Executive Project Status Pack delivery set."),
            ("raid-log-action-tracker-bundle.zip", "Full RAID Log + Action Tracker Bundle delivery set."),
            ("client-kickoff-meeting-pack.zip", "Full Client Kickoff Meeting Pack delivery set."),
            ("website-redesign-planning-kit.zip", "Full Website Redesign Planning Kit delivery set."),
            ("consulting-proposal-starter-kit.zip", "Full Consulting Proposal Starter Kit delivery set."),
            ("bonus-executive-communication-templates.docx", "Bonus communication templates for sharper updates, asks, and readouts."),
        ],
        "docs": {
            "README.md": """
            # Full Clarpoint Business Execution Toolkit

            This bundle gives you the full Clarpoint starter system: project control, leadership communication, client kickoff structure, proposal support, and website planning resources in one package.
            """,
            "bonus-executive-communication-templates.md": """
            # Bonus Executive Communication Templates

            ## Leadership update opener
            The work is moving, but the main point for leadership this week is [insert core message].

            ## Decision request structure
            - Decision needed:
            - Why now:
            - Recommended option:
            - If delayed, impact is:

            ## Risk framing structure
            - Risk:
            - Why it matters:
            - Current mitigation:
            - Help needed:

            ## Meeting readout structure
            - What we aligned on:
            - What changed:
            - What happens next:
            """,
        },
        "csvs": {},
    },
}


def markdown_to_docx(markdown_text: str, destination: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            p = document.add_paragraph()
            p.style = document.styles["Title"]
            run = p.add_run(line[2:].strip())
            run.bold = True
            continue
        if line.startswith("## "):
            p = document.add_paragraph()
            p.style = document.styles["Heading 1"]
            p.add_run(line[3:].strip())
            continue
        if line.startswith("### "):
            p = document.add_paragraph()
            p.style = document.styles["Heading 2"]
            p.add_run(line[4:].strip())
            continue
        if line.startswith("- [ ] "):
            p = document.add_paragraph(style="List Bullet")
            p.add_run("□ " + line[6:].strip())
            continue
        if line.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            p.add_run(line[2:].strip())
            continue
        p = document.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.save(destination)


def write_text(path: Path, content: str) -> None:
    path.write_text(clean(content), encoding="utf-8")


def write_csv(path: Path, rows: list[list[str]]) -> None:
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerows(rows)


def create_pdf_summary(folder: Path, pack: dict) -> None:
    path = folder / "product-summary.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=LETTER, leftMargin=0.7 * inch, rightMargin=0.7 * inch, topMargin=0.7 * inch, bottomMargin=0.7 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ClarpointTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=22, textColor=colors.HexColor("#132322"), spaceAfter=12)
    h_style = ParagraphStyle("ClarpointHeading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, textColor=colors.HexColor("#0f766e"), spaceBefore=8, spaceAfter=6)
    body_style = ParagraphStyle("ClarpointBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=14, textColor=colors.HexColor("#223434"))
    small_style = ParagraphStyle("ClarpointSmall", parent=styles["BodyText"], fontName="Helvetica", fontSize=9, leading=12, textColor=colors.HexColor("#4b5b5a"))

    story = [
        Paragraph(pack["title"], title_style),
        Paragraph(pack["description"], body_style),
        Spacer(1, 0.12 * inch),
        Paragraph(BRAND["tagline"], small_style),
        Spacer(1, 0.18 * inch),
        Paragraph("Who this pack is for", h_style),
        Paragraph(pack["audience"], body_style),
        Paragraph("What the buyer receives", h_style),
    ]

    bullet_items = [ListItem(Paragraph(f"<b>{name}</b> — {desc}", body_style), leftIndent=8) for name, desc in pack["deliverables"]]
    story.append(ListFlowable(bullet_items, bulletType="bullet", start="circle", leftIndent=14))
    story.append(Paragraph("What this pack should help improve", h_style))
    outcome_items = [ListItem(Paragraph(item, body_style), leftIndent=8) for item in pack["outcomes"]]
    story.append(ListFlowable(outcome_items, bulletType="bullet", start="square", leftIndent=14))
    story.append(Paragraph("Recommended use order", h_style))
    workflow_items = [ListItem(Paragraph(item, body_style), leftIndent=8) for item in pack["workflow"]]
    story.append(ListFlowable(workflow_items, bulletType="1", leftIndent=14))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph("Need this customized? Clarpoint can adapt these tools into a fuller operating model, website plan, executive readout, or project delivery system.", small_style))
    doc.build(story)


def build_pack_guide_markdown(pack: dict) -> str:
    lines = [
        f"# {pack['title']} Guide",
        "",
        pack["description"],
        "",
        "## Who this pack is for",
        pack["audience"],
        "",
        "## What the buyer receives",
    ]
    for name, desc in pack["deliverables"]:
        lines.append(f"- {name}: {desc}")
    lines.extend(["", "## What this pack should help improve"])
    for item in pack["outcomes"]:
        lines.append(f"- {item}")
    lines.extend(["", "## Recommended use order"])
    for idx, item in enumerate(pack["workflow"], start=1):
        lines.append(f"{idx}. {item}")
    lines.extend(["", "## Support note", f"If you want this customized for your business, {BRAND['name']} can adapt the toolkit into a fuller delivery, communication, or website support model."])
    return "\n".join(lines)


def write_included_files(folder: Path, pack: dict) -> None:
    lines = [pack["title"], "", "Included files in this folder:"]
    for filename, desc in pack["deliverables"]:
        lines.append(f"- {filename}: {desc}")
    lines.extend(["", "Also included:", "- README.md", "- how-to-use.txt", "- included-files.txt", "- product-summary.pdf", f"- {pack['guide_filename']}"])
    (folder / "included-files.txt").write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def write_how_to_use(folder: Path, pack: dict) -> None:
    lines = [pack["title"], "", "Recommended use order:"]
    for idx, step in enumerate(pack["workflow"], start=1):
        lines.append(f"{idx}. {step}")
    lines.extend(["", BRAND["tagline"]])
    (folder / "how-to-use.txt").write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def build_docs() -> None:
    for slug, pack in PACKS.items():
        folder = DOWNLOADS / slug
        folder.mkdir(parents=True, exist_ok=True)

        for filename, content in pack["docs"].items():
            target = folder / filename
            write_text(target, content)
            if filename.endswith(".md") and filename != "README.md":
                markdown_to_docx(clean(content), folder / f"{target.stem}.docx")

        guide_markdown = build_pack_guide_markdown(pack)
        guide_md_path = folder / f"{Path(pack['guide_filename']).stem}.md"
        write_text(guide_md_path, guide_markdown)
        markdown_to_docx(guide_markdown, folder / pack["guide_filename"])

        if slug == "full-clarpoint-business-execution-toolkit":
            bonus_markdown = clean(pack["docs"]["bonus-executive-communication-templates.md"])
            markdown_to_docx(bonus_markdown, folder / "bonus-executive-communication-templates.docx")

        for filename, rows in pack["csvs"].items():
            write_csv(folder / filename, rows)

        write_included_files(folder, pack)
        write_how_to_use(folder, pack)
        create_pdf_summary(folder, pack)


if __name__ == "__main__":
    build_docs()
